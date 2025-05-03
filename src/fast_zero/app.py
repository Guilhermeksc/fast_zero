import os
from io import BytesIO
from zipfile import ZipFile, ZIP_DEFLATED
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
import psycopg2
from psycopg2.extras import RealDictCursor
from docx import Document

load_dotenv()

def get_conn():
    return psycopg2.connect(
        host=os.getenv("DB__HOST"),
        port=int(os.getenv("DB__PORT", 5432)),
        user=os.getenv("DB__USERNAME"),
        password=os.getenv("DB__PASSWORD"),
        dbname=os.getenv("DB__NAME"),
        cursor_factory=RealDictCursor
    )

app = FastAPI()

@app.get(
    "/pncp/{sequencial}/{ano}/zip",
    response_class=StreamingResponse,
    summary="Baixar ZIP com DOCX dos dados PNCP"
)
def download_pncp_zip(sequencial: int, ano: int):
    # 1) consulta
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            "SELECT * FROM pncp WHERE sequencial_compra=%s AND ano_compra=%s",
            (sequencial, ano)
        )
        row = cur.fetchone()
    if not row:
        raise HTTPException(404, "Registro não encontrado")

    # 2) gera .docx em memória
    docx_buf = BytesIO()
    doc = Document()
    doc.add_heading(f"PNCP {sequencial}/{ano}", level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Campo", "Valor"
    for k, v in row.items():
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = "" if v is None else str(v)
    doc.save(docx_buf)
    docx_buf.seek(0)

    # 3) empacota num ZIP
    zip_buf = BytesIO()
    with ZipFile(zip_buf, "w", ZIP_DEFLATED) as z:
        z.writestr(f"pncp_{sequencial}_{ano}.docx", docx_buf.getvalue())
    zip_buf.seek(0)

    # 4) retorna como download
    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={
            "Content-Disposition": f"attachment; filename=pncp_{sequencial}_{ano}.zip"
        }
    )
